#!/usr/bin/env python3
"""Fill the HLS Bridge Advisory CV template with structured candidate data.

This script does no PDF reading, no content judgment, and no PII decisions --
it only takes a JSON data file (already stripped of PII, already reviewed
for the Executive Summary) and mechanically fills the branded Word template,
matching the exact styles already shipped in Formatted CVs/.

Usage:
    python3 tools/generate_cv.py --data candidate.json --photo photo.jpg

See --help for all options.
"""
import argparse
import copy
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

HEADINGS_IN_ORDER = [
    "EXECUTIVE SUMMARY",
    "CORE SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "EDUCATION",
    "CERTIFICATIONS & AWARDS",
]

REQUIRED_TOP_LEVEL = ["full_name", "title"]
LIST_FIELDS_DEFAULT = ["skills", "experience", "education", "certifications"]


def validate_data(data):
    if not isinstance(data, dict):
        raise ValueError("Candidate data must be a JSON object.")

    for key in REQUIRED_TOP_LEVEL:
        if not data.get(key):
            raise ValueError(f"Candidate data is missing required field: {key!r}")

    data.setdefault("executive_summary", "")
    for key in LIST_FIELDS_DEFAULT:
        data.setdefault(key, [])
        if not isinstance(data[key], list):
            raise ValueError(f"{key!r} must be a list.")

    for i, job in enumerate(data["experience"]):
        for key in ("company_role", "dates"):
            if not job.get(key):
                raise ValueError(f"experience[{i}] is missing required field: {key!r}")
        job.setdefault("bullets", [])
        if not isinstance(job["bullets"], list):
            raise ValueError(f"experience[{i}]['bullets'] must be a list.")

    for i, edu in enumerate(data["education"]):
        for key in ("degree_line", "detail_line"):
            if not edu.get(key):
                raise ValueError(f"education[{i}] is missing required field: {key!r}")

    return data


def find_heading_positions(all_paragraphs):
    positions = {}
    for text in HEADINGS_IN_ORDER:
        idx = next(
            (i for i, p in enumerate(all_paragraphs) if p.text.strip() == text),
            None,
        )
        if idx is None:
            raise RuntimeError(
                f"Template is missing expected heading: {text!r}. "
                "Has HLS_CV_Template.docx been edited?"
            )
        positions[text] = idx
    return positions


def section_slice(all_paragraphs, positions, heading_text):
    idx = positions[heading_text]
    order = HEADINGS_IN_ORDER.index(heading_text)
    if order + 1 < len(HEADINGS_IN_ORDER):
        next_text = HEADINGS_IN_ORDER[order + 1]
        end_idx = positions[next_text]
    else:
        end_idx = len(all_paragraphs)
    return all_paragraphs[idx + 1:end_idx]


def clone_paragraph(template_element, parent):
    new_element = copy.deepcopy(template_element)
    return new_element, Paragraph(new_element, parent)


def remove_paragraphs(paragraphs):
    for p in paragraphs:
        p._element.getparent().remove(p._element)


def set_run_text(paragraph, run_index, text):
    paragraph.runs[run_index].text = text


def fill_experience(doc, all_paragraphs, positions, experience):
    heading = all_paragraphs[positions["PROFESSIONAL EXPERIENCE"]]
    section = section_slice(all_paragraphs, positions, "PROFESSIONAL EXPERIENCE")

    header_tpl = next(p for p in section if p.style.name == "Normal")
    bullet_tpl = next(p for p in section if p.style.name == "List Bullet")
    header_tpl_el = copy.deepcopy(header_tpl._element)
    bullet_tpl_el = copy.deepcopy(bullet_tpl._element)
    parent = heading._parent

    remove_paragraphs(section)

    anchor = heading._element
    for job in experience:
        header_el, header_p = clone_paragraph(header_tpl_el, parent)
        anchor.addnext(header_el)
        anchor = header_el
        set_run_text(header_p, 0, job["company_role"])
        set_run_text(header_p, -1, job["dates"])

        for bullet_text in job["bullets"]:
            bullet_el, bullet_p = clone_paragraph(bullet_tpl_el, parent)
            anchor.addnext(bullet_el)
            anchor = bullet_el
            set_run_text(bullet_p, 0, bullet_text)


def fill_education(doc, all_paragraphs, positions, education):
    heading = all_paragraphs[positions["EDUCATION"]]
    section = section_slice(all_paragraphs, positions, "EDUCATION")

    degree_tpl = next(p for p in section if p.runs and p.runs[0].font.bold)
    detail_tpl = next(
        p for p in section if p.runs and p.runs[0].font.italic and not p.runs[0].font.bold
    )
    degree_tpl_el = copy.deepcopy(degree_tpl._element)
    detail_tpl_el = copy.deepcopy(detail_tpl._element)
    parent = heading._parent

    remove_paragraphs(section)

    anchor = heading._element
    for edu in education:
        degree_el, degree_p = clone_paragraph(degree_tpl_el, parent)
        anchor.addnext(degree_el)
        anchor = degree_el
        set_run_text(degree_p, 0, edu["degree_line"])

        detail_el, detail_p = clone_paragraph(detail_tpl_el, parent)
        anchor.addnext(detail_el)
        anchor = detail_el
        set_run_text(detail_p, 0, edu["detail_line"])


def fill_certifications(doc, all_paragraphs, positions, certifications):
    heading = all_paragraphs[positions["CERTIFICATIONS & AWARDS"]]
    section = section_slice(all_paragraphs, positions, "CERTIFICATIONS & AWARDS")

    if not certifications:
        remove_paragraphs(section)
        heading._element.getparent().remove(heading._element)
        return

    bullet_tpl = next(p for p in section if p.style.name == "List Bullet")
    bullet_tpl_el = copy.deepcopy(bullet_tpl._element)
    parent = heading._parent

    remove_paragraphs(section)

    anchor = heading._element
    for cert_text in certifications:
        bullet_el, bullet_p = clone_paragraph(bullet_tpl_el, parent)
        anchor.addnext(bullet_el)
        anchor = bullet_el
        set_run_text(bullet_p, 0, cert_text)


def fill_executive_summary(all_paragraphs, positions, summary_text):
    para = all_paragraphs[positions["EXECUTIVE SUMMARY"] + 1]
    text = summary_text.strip()
    if text:
        set_run_text(para, 0, text)
    else:
        set_run_text(para, 0, " ")
        para.paragraph_format.space_after = Pt(14)


def fill_core_skills(all_paragraphs, positions, skills):
    para = all_paragraphs[positions["CORE SKILLS"] + 1]
    set_run_text(para, 0, "  •  ".join(skills))


def fill_name_title_photo(doc, full_name, title, photo_path):
    table = doc.tables[0]
    cell0 = table.rows[0].cells[0]
    set_run_text(cell0.paragraphs[0], 0, full_name)
    set_run_text(cell0.paragraphs[1], 0, title)

    if photo_path:
        cell1 = table.rows[0].cells[1]
        para = cell1.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        run.add_picture(photo_path, width=Inches(1.15), height=Inches(1.15))


def build_output_path(full_name, output_dir, output_override, date_str):
    if output_override:
        return output_override
    safe_name = full_name.replace("/", "-").replace("\\", "-")
    filename = f"{safe_name}_HLS_{date_str}.docx"
    return os.path.join(output_dir, filename)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data", required=True, help="Path to candidate JSON data file")
    parser.add_argument("--photo", help="Path to a pre-cropped square candidate photo")
    parser.add_argument(
        "--template",
        default=os.path.join("Formatted CVs", "HLS_CV_Template.docx"),
        help="Path to the blank HLS CV template",
    )
    parser.add_argument(
        "--output-dir",
        default="Formatted CVs",
        help="Directory for the output file (used only if --output is omitted)",
    )
    parser.add_argument("--output", help="Explicit output .docx path (overrides --output-dir)")
    parser.add_argument(
        "--date", default=None, help="Date suffix in ddmmyyyy form (default: today)"
    )
    parser.add_argument(
        "--overwrite", action="store_true", help="Allow overwriting an existing output file"
    )
    args = parser.parse_args()

    if not os.path.isfile(args.data):
        print(f"Error: data file not found: {args.data}", file=sys.stderr)
        sys.exit(1)
    if not os.path.isfile(args.template):
        print(f"Error: template not found: {args.template}", file=sys.stderr)
        sys.exit(1)
    if args.photo and not os.path.isfile(args.photo):
        print(f"Error: photo not found: {args.photo}", file=sys.stderr)
        sys.exit(1)

    with open(args.data, "r", encoding="utf-8") as f:
        raw = json.load(f)

    try:
        data = validate_data(raw)
    except ValueError as e:
        print(f"Error: invalid candidate data -- {e}", file=sys.stderr)
        sys.exit(1)

    date_str = args.date or datetime.now().strftime("%d%m%Y")
    output_path = build_output_path(data["full_name"], args.output_dir, args.output, date_str)

    if os.path.exists(output_path) and not args.overwrite:
        print(
            f"Error: output already exists: {output_path} (pass --overwrite to replace it)",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document(args.template)
    all_paragraphs = list(doc.paragraphs)
    positions = find_heading_positions(all_paragraphs)

    fill_name_title_photo(doc, data["full_name"], data["title"], args.photo)
    fill_executive_summary(all_paragraphs, positions, data["executive_summary"])
    fill_core_skills(all_paragraphs, positions, data["skills"])
    fill_experience(doc, all_paragraphs, positions, data["experience"])
    fill_education(doc, all_paragraphs, positions, data["education"])
    fill_certifications(doc, all_paragraphs, positions, data["certifications"])

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)

    print(os.path.abspath(output_path))


if __name__ == "__main__":
    main()
